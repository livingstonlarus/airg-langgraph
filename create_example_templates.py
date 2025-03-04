#!/usr/bin/env python3
"""
AIRG-LangGraph - Create Example Templates
Creates example resume and cover letter templates
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_resume_template():
    """
    Create an example resume template
    """
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Add name
    name = doc.add_paragraph()
    name_run = name.add_run("Your Name")
    name_run.bold = True
    name_run.font.size = Pt(18)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add contact info
    contact = doc.add_paragraph()
    contact.add_run("email@example.com | (123) 456-7890 | City, State | LinkedIn: linkedin.com/in/yourname")
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add professional summary
    doc.add_heading("Professional Summary", level=1)
    summary = doc.add_paragraph()
    summary.add_run("{{SUMMARY}}")
    
    # Add skills
    doc.add_heading("Skills", level=1)
    skills = doc.add_paragraph()
    skills.add_run("{{SKILLS}}")
    
    # Add experience
    doc.add_heading("Experience", level=1)
    experience = doc.add_paragraph()
    experience.add_run("{{EXPERIENCE}}")
    
    # Add education
    doc.add_heading("Education", level=1)
    education = doc.add_paragraph()
    education.add_run("{{EDUCATION}}")
    
    # Add projects
    doc.add_heading("Projects", level=1)
    projects = doc.add_paragraph()
    projects.add_run("{{PROJECTS}}")
    
    # Add certifications
    doc.add_heading("Certifications", level=1)
    certifications = doc.add_paragraph()
    certifications.add_run("{{CERTIFICATIONS}}")
    
    # Create examples directory if it doesn't exist
    os.makedirs("examples", exist_ok=True)
    
    # Save the document
    doc.save("examples/resume_template.docx")
    print("Resume template created: examples/resume_template.docx")


def create_cover_letter_template():
    """
    Create an example cover letter template
    """
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add date
    date = doc.add_paragraph()
    date.add_run("Date: [Current Date]")
    
    # Add recipient
    recipient = doc.add_paragraph()
    recipient.add_run("{{GREETING}}")
    
    # Add introduction
    introduction = doc.add_paragraph()
    introduction.add_run("{{INTRODUCTION}}")
    
    # Add body
    body = doc.add_paragraph()
    body.add_run("{{BODY}}")
    
    # Add closing
    closing = doc.add_paragraph()
    closing.add_run("{{CLOSING}}")
    
    # Add signature
    signature = doc.add_paragraph()
    signature.add_run("{{SIGNATURE}}")
    signature.add_run("\nYour Name\nemail@example.com\n(123) 456-7890")
    
    # Create examples directory if it doesn't exist
    os.makedirs("examples", exist_ok=True)
    
    # Save the document
    doc.save("examples/cover_letter_template.docx")
    print("Cover letter template created: examples/cover_letter_template.docx")


if __name__ == "__main__":
    create_resume_template()
    create_cover_letter_template()
    print("\nExample templates created successfully!")
    print("You can use these templates with the AIRG-LangGraph application:")
    print("python main.py --resume-template examples/resume_template.docx --cover-letter-template examples/cover_letter_template.docx ...")