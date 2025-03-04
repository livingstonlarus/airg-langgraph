"""
AIRG-LangGraph - Utilities for generating PDF files
"""

import os
import tempfile
import subprocess
from typing import Optional
from weasyprint import HTML


def docx_to_html(docx_path: str) -> str:
    """
    Convert a DOCX file to HTML using pandoc
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Path to the generated HTML file
    """
    # Create a temporary file for the HTML output
    fd, html_path = tempfile.mkstemp(suffix=".html")
    os.close(fd)
    
    # Use pandoc to convert DOCX to HTML
    try:
        subprocess.run(
            ["pandoc", docx_path, "-o", html_path],
            check=True,
            capture_output=True,
        )
    except (subprocess.CalledProcessError, FileNotFoundError) as e:
        # If pandoc is not installed or fails, use a simpler approach
        print(f"Warning: Pandoc conversion failed ({str(e)}). Using fallback method.")
        return _fallback_docx_to_html(docx_path)
    
    return html_path


def _fallback_docx_to_html(docx_path: str) -> str:
    """
    Fallback method to convert DOCX to HTML using python-docx
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Path to the generated HTML file
    """
    from docx import Document
    
    # Load the document
    doc = Document(docx_path)
    
    # Create a simple HTML representation
    html_content = "<html><head><style>body { font-family: Arial, sans-serif; }</style></head><body>"
    
    # Add paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            html_content += f"<p>{paragraph.text}</p>"
    
    # Add tables
    for table in doc.tables:
        html_content += "<table border='1' cellpadding='3'>"
        for row in table.rows:
            html_content += "<tr>"
            for cell in row.cells:
                cell_text = ""
                for paragraph in cell.paragraphs:
                    cell_text += paragraph.text + "<br/>"
                html_content += f"<td>{cell_text}</td>"
            html_content += "</tr>"
        html_content += "</table>"
    
    html_content += "</body></html>"
    
    # Write to a temporary file
    fd, html_path = tempfile.mkstemp(suffix=".html")
    with os.fdopen(fd, 'w') as f:
        f.write(html_content)
    
    return html_path


def html_to_pdf(html_path: str, pdf_path: str) -> str:
    """
    Convert an HTML file to PDF using WeasyPrint
    
    Args:
        html_path: Path to the HTML file
        pdf_path: Path to save the PDF file
        
    Returns:
        Path to the generated PDF file
    """
    # Create the directory if it doesn't exist
    os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
    
    # Convert HTML to PDF
    HTML(html_path).write_pdf(pdf_path)
    
    return pdf_path


def docx_to_pdf(docx_path: str, pdf_path: str) -> str:
    """
    Convert a DOCX file to PDF
    
    Args:
        docx_path: Path to the DOCX file
        pdf_path: Path to save the PDF file
        
    Returns:
        Path to the generated PDF file
    """
    # Convert DOCX to HTML
    html_path = docx_to_html(docx_path)
    
    try:
        # Convert HTML to PDF
        pdf_path = html_to_pdf(html_path, pdf_path)
    finally:
        # Clean up the temporary HTML file
        if os.path.exists(html_path):
            os.remove(html_path)
    
    return pdf_path