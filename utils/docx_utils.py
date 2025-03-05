"""
AIRG-LangGraph - Utilities for working with DOCX files
"""

import re
import os
from typing import Dict, List, Tuple, Any
from docx import Document


def read_document(document_path: str) -> Tuple[Document, str, Dict[str, List[str]]]:
    """
    Read a DOCX document and analyze its content
    
    Args:
        document_path: Path to the DOCX document
        
    Returns:
        Tuple containing the Document object, text content, and section content
    """
    # Load the document
    doc = Document(document_path)
    
    # Extract text from the document
    text_content = ""
    for paragraph in doc.paragraphs:
        text_content += paragraph.text + "\n"
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text_content += paragraph.text + "\n"
    
    # Analyze document structure to identify sections
    sections = analyze_document_structure(doc)
    
    return doc, text_content, sections


def analyze_document_structure(doc: Document) -> Dict[str, List[str]]:
    """
    Analyze the document structure to identify sections
    
    Args:
        doc: Document object to analyze
        
    Returns:
        Dictionary mapping section names to their content
    """
    sections = {
        "personal_info": [],
        "summary": [],
        "experience": [],
        "skills": [],
        "education": [],
        "other": []
    }
    
    current_section = "other"
    
    # Simple heuristic to identify sections based on heading styles and content
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        # Skip empty paragraphs
        if not text:
            continue
        
        # Check if this is a section heading
        if paragraph.style.name.startswith('Heading') or (text.isupper() and len(text) < 30):
            lower_text = text.lower()
            
            if any(keyword in lower_text for keyword in ["profile", "summary", "objective", "about"]):
                current_section = "summary"
            elif any(keyword in lower_text for keyword in ["experience", "employment", "work", "career"]):
                current_section = "experience"
            elif any(keyword in lower_text for keyword in ["skill", "expertise", "competenc", "proficienc"]):
                current_section = "skills"
            elif any(keyword in lower_text for keyword in ["education", "academic", "qualification", "degree"]):
                current_section = "education"
            elif any(keyword in lower_text for keyword in ["contact", "personal", "info"]):
                current_section = "personal_info"
            else:
                current_section = "other"
        
        # Add the paragraph text to the current section
        sections[current_section].append(text)
    
    return sections


def update_document_content(doc: Document, updated_sections: Dict[str, List[str]]) -> Document:
    """
    Update a DOCX document with the provided section content
    
    Args:
        doc: Document object to update
        updated_sections: Dictionary mapping section names to their updated content
        
    Returns:
        Updated Document object
    """
    # Create a new document to avoid modifying the original
    new_doc = Document()
    
    # Copy styles from the original document
    for style in doc.styles:
        if style.name not in new_doc.styles:
            try:
                new_doc.styles.add_style(style.name, style.type)
            except:
                pass  # Skip if style already exists or can't be added
    
    # Analyze the original document structure
    original_sections = analyze_document_structure(doc)
    
    # Track which sections we've processed
    processed_sections = set()
    
    # Copy the document paragraph by paragraph, replacing section content as needed
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        
        # Determine which section this paragraph belongs to
        current_section = None
        for section_name, section_content in original_sections.items():
            if text in section_content and section_name not in processed_sections:
                current_section = section_name
                break
        
        # If this is a section we need to update and haven't processed yet
        if current_section and current_section in updated_sections and current_section not in processed_sections:
            # Add the updated content for this section
            for updated_text in updated_sections[current_section]:
                p = new_doc.add_paragraph(updated_text)
                # Try to copy the style
                try:
                    p.style = paragraph.style.name
                except:
                    pass  # Skip if style can't be applied
            
            # Mark this section as processed
            processed_sections.add(current_section)
        else:
            # Copy the paragraph as is
            p = new_doc.add_paragraph(paragraph.text)
            # Try to copy the style
            try:
                p.style = paragraph.style.name
            except:
                pass  # Skip if style can't be applied
    
    # Copy tables
    for table in doc.tables:
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        
        # Copy table style
        try:
            new_table.style = table.style
        except:
            pass  # Skip if style can't be applied
        
        # Copy cell content
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    new_cell = new_table.cell(i, j)
                    new_cell.text = paragraph.text
    
    return new_doc


def save_document(doc: Document, output_path: str) -> str:
    """
    Save a Document object to a file
    
    Args:
        doc: Document object to save
        output_path: Path to save the document to
        
    Returns:
        Path to the saved document
    """
    # Create the directory if it doesn't exist
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Save the document
    doc.save(output_path)
    
    return output_path


def process_template(template_path: str) -> Dict[str, Any]:
    """
    Process a DOCX template and analyze its content
    
    Args:
        template_path: Path to the DOCX template
        
    Returns:
        Dictionary containing the Document object, text content, and section content
    """
    # Read and analyze the document
    doc, text_content, sections = read_document(template_path)
    
    return {
        "document": doc,
        "text_content": text_content,
        "sections": sections,
    }